from dataclasses import dataclass
from io import BytesIO

import fitz
from docx import Document
from fastapi import HTTPException


@dataclass
class ParsedDocument:
    filename: str
    raw_text: str
    char_count: int
    parse_quality: str = "good"


def _decode_txt(content: bytes) -> str:
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore").strip()


def _parse_pdf(content: bytes) -> tuple[str, str]:
    doc = fitz.open(stream=content, filetype="pdf")
    try:
        parts = [page.get_text() for page in doc]
        text = "\n".join(parts).strip()
    finally:
        doc.close()
    if not text:
        return "", "scanned"
    if len(text) < 100:
        return text, "low"
    return text, "good"


def _parse_docx(content: bytes) -> tuple[str, str]:
    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        return "", "low"
    if len(text) < 100:
        return text, "low"
    return text, "good"


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    if not content:
        raise HTTPException(status_code=422, detail=f"File '{filename}' is empty.")

    lower = filename.lower()
    parse_quality = "good"
    try:
        if lower.endswith(".pdf"):
            text, parse_quality = _parse_pdf(content)
            if parse_quality == "scanned":
                raise HTTPException(
                    status_code=422,
                    detail=(
                        f"File '{filename}' appears to be a scanned PDF with no extractable text. "
                        "Please upload a searchable PDF or DOCX."
                    ),
                )
        elif lower.endswith(".docx"):
            text, parse_quality = _parse_docx(content)
        elif lower.endswith(".doc"):
            raise HTTPException(
                status_code=422,
                detail=f"File '{filename}': .doc format not supported. Please convert to .docx or PDF.",
            )
        elif lower.endswith(".txt"):
            text = _decode_txt(content)
            parse_quality = "low" if len(text) < 100 else "good"
        else:
            raise HTTPException(
                status_code=422,
                detail=f"File '{filename}': unsupported format. Use PDF, DOCX, or TXT.",
            )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=422, detail=f"Failed to parse '{filename}': {exc}"
        ) from exc

    if not text:
        raise HTTPException(
            status_code=422,
            detail=f"File '{filename}' contains no extractable text.",
        )

    return ParsedDocument(
        filename=filename,
        raw_text=text,
        char_count=len(text),
        parse_quality=parse_quality,
    )
